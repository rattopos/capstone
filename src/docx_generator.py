"""
DOCX 생성 모듈
10개 템플릿을 순서대로 처리하여 DOCX로 생성
정답 이미지와 유사하게 포맷팅
"""

from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, Inches, Twips
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from src.base_generator import BaseDocumentGenerator, TEMPLATE_ORDER


class DOCXGenerator(BaseDocumentGenerator):
    """DOCX 생성 클래스"""
    
    # 시도 이름 포맷팅 (정답 이미지와 동일하게 글자 사이에 공백 추가)
    REGION_FORMAT = {
        '전국': '전 국',
        '서울': '서 울',
        '부산': '부 산',
        '대구': '대 구',
        '인천': '인 천',
        '광주': '광 주',
        '대전': '대 전',
        '울산': '울 산',
        '세종': '세 종',
        '경기': '경 기',
        '강원': '강 원',
        '충북': '충 북',
        '충남': '충 남',
        '전북': '전 북',
        '전남': '전 남',
        '경북': '경 북',
        '경남': '경 남',
        '제주': '제 주',
    }
    
    def check_docx_generator_available(self):
        """DOCX 생성 라이브러리 사용 가능 여부 확인"""
        try:
            from docx import Document
            return True, None
        except ImportError:
            return False, (
                'DOCX 생성 라이브러리가 설치되지 않았습니다. 다음을 설치해주세요:\n'
                'pip install python-docx'
            )
    
    def _format_region_name(self, text):
        """시도 이름을 정답 형식으로 포맷팅"""
        text = text.strip()
        return self.REGION_FORMAT.get(text, text)
    
    def _set_cell_border(self, cell, border_size=4):
        """셀 테두리 설정"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{border_size}" w:space="0" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="{border_size}" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="{border_size}" w:space="0" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="{border_size}" w:space="0" w:color="000000"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)
    
    def _set_cell_shading(self, cell, color="F5F5F5"):
        """셀 배경색 설정"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        tcPr.append(shading)
    
    def html_to_docx_paragraph(self, doc, html_content):
        """HTML 내용을 DOCX 단락으로 변환"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 스타일 태그 제거
        for style in soup.find_all(['style', 'script']):
            style.decompose()
        
        body = soup.find('body') or soup
        
        for element in body.children:
            self._process_element(doc, element)
    
    def _process_element(self, doc, element):
        """단일 HTML 요소를 DOCX로 변환"""
        if not hasattr(element, 'name') or element.name is None:
            return
        
        tag_name = element.name.lower()
        text = element.get_text(strip=True)
        
        # 헤딩 처리 - 스타일 적용
        if tag_name in ('h1', 'h2', 'h3', 'h4'):
            level = int(tag_name[1])
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(14 - level)  # h1=13pt, h2=12pt, etc.
                para.space_after = Pt(6)
        
        # section-title 또는 subsection-title 처리
        elif tag_name == 'div':
            classes = element.get('class', [])
            
            if 'section-title' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(12)
                    para.space_before = Pt(12)
                    para.space_after = Pt(6)
            
            elif 'subsection-title' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(11)
                    para.space_before = Pt(10)
                    para.space_after = Pt(4)
            
            elif 'content-text' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(10)
                    para.space_after = Pt(4)
            
            elif 'key-section' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(10)
                    para.space_before = Pt(8)
                    para.space_after = Pt(4)
            
            elif 'key-item' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(9)
                    para.space_after = Pt(2)
            
            elif 'table-title' in classes:
                if text:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(10)
                    para.space_before = Pt(12)
                    para.space_after = Pt(2)
            
            elif 'table-subtitle' in classes:
                if text:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = para.add_run(text)
                    run.font.size = Pt(8)
                    para.space_after = Pt(4)
            
            elif 'footnote' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(8)
                    para.space_before = Pt(4)
            
            elif 'source' in classes:
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.size = Pt(8)
                    para.space_before = Pt(2)
            
            else:
                # 다른 div는 자식 요소 처리
                for child in element.children:
                    self._process_element(doc, child)
        
        # 문단 처리
        elif tag_name == 'p':
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(10)
                para.space_after = Pt(4)
        
        # 테이블 처리
        elif tag_name == 'table':
            self._html_table_to_docx(doc, element)
        
        # 줄바꿈
        elif tag_name in ('br', 'hr'):
            doc.add_paragraph()
        
        # 리스트 처리
        elif tag_name == 'ul':
            for li in element.find_all('li', recursive=False):
                li_text = li.get_text(strip=True)
                if li_text:
                    para = doc.add_paragraph(li_text, style='List Bullet')
                    for run in para.runs:
                        run.font.size = Pt(9)
        elif tag_name == 'ol':
            for li in element.find_all('li', recursive=False):
                li_text = li.get_text(strip=True)
                if li_text:
                    para = doc.add_paragraph(li_text, style='List Number')
                    for run in para.runs:
                        run.font.size = Pt(9)
        
        # 기타 요소
        elif text:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(10)
    
    def _html_table_to_docx(self, doc, table_element):
        """HTML 테이블을 DOCX 테이블로 변환 (정답 이미지 형식)"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # 열 수 계산 (colspan 고려)
        max_cols = 0
        for row in rows:
            col_count = 0
            for cell in row.find_all(['th', 'td']):
                colspan = int(cell.get('colspan', 1))
                col_count += colspan
            max_cols = max(max_cols, col_count)
        
        if max_cols == 0:
            return
        
        # 테이블 생성 (행 수는 나중에 조정)
        table = doc.add_table(rows=0, cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 전체 너비 설정 (페이지 너비에 맞춤)
        table.autofit = False
        
        # 각 열의 너비 계산 (첫 번째 열은 좁게, 나머지는 균등)
        first_col_width = Cm(1.5)
        remaining_width = Cm(15.5)  # 전체 너비 약 17cm
        other_col_width = remaining_width / (max_cols - 1) if max_cols > 1 else remaining_width
        
        # rowspan 추적을 위한 배열
        rowspan_tracker = [0] * max_cols
        
        for row_idx, row in enumerate(rows):
            # 새 행 추가
            table_row = table.add_row()
            cells = row.find_all(['th', 'td'])
            
            col_idx = 0
            cell_idx = 0
            
            while col_idx < max_cols and cell_idx < len(cells):
                # rowspan이 진행 중인 셀 건너뛰기
                while col_idx < max_cols and rowspan_tracker[col_idx] > 0:
                    rowspan_tracker[col_idx] -= 1
                    col_idx += 1
                
                if col_idx >= max_cols:
                    break
                
                cell = cells[cell_idx]
                cell_text = cell.get_text(strip=True)
                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))
                
                # 첫 번째 열 셀인 경우 시도 이름 포맷팅
                if col_idx == 0:
                    cell_text = self._format_region_name(cell_text)
                
                # 셀에 텍스트 설정
                docx_cell = table_row.cells[col_idx]
                docx_cell.text = cell_text
                
                # 셀 정렬
                para = docx_cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 폰트 크기 설정 (작은 폰트로 한 페이지에 맞춤)
                for run in para.runs:
                    run.font.size = Pt(8)
                    if cell.name == 'th':
                        run.bold = True
                
                # 헤더 셀 배경색
                if cell.name == 'th':
                    self._set_cell_shading(docx_cell, "F5F5F5")
                
                # 테두리 설정
                self._set_cell_border(docx_cell)
                
                # colspan 처리 (셀 병합)
                if colspan > 1:
                    for i in range(1, colspan):
                        if col_idx + i < max_cols:
                            merge_cell = table_row.cells[col_idx + i]
                            docx_cell.merge(merge_cell)
                            self._set_cell_border(merge_cell)
                
                # rowspan 추적 설정
                if rowspan > 1:
                    for i in range(colspan):
                        if col_idx + i < max_cols:
                            rowspan_tracker[col_idx + i] = rowspan - 1
                
                col_idx += colspan
                cell_idx += 1
            
            # 남은 rowspan 감소
            while col_idx < max_cols:
                if rowspan_tracker[col_idx] > 0:
                    rowspan_tracker[col_idx] -= 1
                col_idx += 1
        
        # 열 너비 설정
        for col_idx, column in enumerate(table.columns):
            if col_idx == 0:
                for cell in column.cells:
                    cell.width = first_col_width
            else:
                for cell in column.cells:
                    cell.width = other_col_width
        
        # 행 높이 설정 (작게)
        for row in table.rows:
            row.height = Pt(16)
        
        # 테이블 후 약간의 여백
        doc.add_paragraph()
    
    def _setup_document_style(self, doc):
        """문서 스타일 설정 (한 페이지에 맞추기 위한 설정)"""
        # 페이지 설정
        for section in doc.sections:
            # 여백 설정 (좁게)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            
            # A4 크기
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    
    def generate_docx(self, excel_path, year, quarter, templates_dir='templates', missing_values=None):
        """
        여러 템플릿을 처리하여 하나의 DOCX 파일로 생성
        
        Args:
            excel_path: 엑셀 파일 경로
            year: 연도
            quarter: 분기
            templates_dir: 템플릿 디렉토리 경로
            missing_values: 사용자가 입력한 결측치 값
            
        Returns:
            tuple: (성공 여부, 결과 dict 또는 에러 메시지)
        """
        # DOCX 생성 라이브러리 확인
        is_available, error_msg = self.check_docx_generator_available()
        if not is_available:
            return False, error_msg
        
        excel_path = Path(excel_path)
        
        # 엑셀 파일 유효성 검증
        is_valid, error_msg = self.validate_excel_file(excel_path)
        if not is_valid:
            return False, error_msg
        
        # 엑셀 추출기 준비
        excel_extractor, flexible_mapper, error_msg = self.prepare_excel_extractor(
            excel_path, year, quarter
        )
        if excel_extractor is None:
            return False, error_msg
        
        try:
            # DOCX 문서 생성
            doc = Document()
            
            # 문서 스타일 설정
            self._setup_document_style(doc)
            
            # 문서 제목
            title = doc.add_heading(f'{year}년 {quarter}분기 지역경제동향 보도자료', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 템플릿 처리
            processed_count = 0
            errors = []
            templates_dir_path = Path(templates_dir)
            
            for template_name in TEMPLATE_ORDER:
                try:
                    result = self._process_single_template(
                        template_name=template_name,
                        templates_dir_path=templates_dir_path,
                        excel_extractor=excel_extractor,
                        flexible_mapper=flexible_mapper,
                        year=year,
                        quarter=quarter,
                        missing_value_overrides=missing_values
                    )
                    
                    if result is None or isinstance(result, str):
                        errors.append(result or f'{template_name}: 템플릿 처리 실패')
                        continue
                    
                    template_name_out, template_html, _ = result
                    
                    # 페이지 구분선 추가 (첫 번째 이후)
                    if processed_count > 0:
                        doc.add_page_break()
                    
                    # HTML 내용을 DOCX로 변환 (섹션 제목은 HTML에 포함되어 있음)
                    self.html_to_docx_paragraph(doc, template_html)
                    
                    processed_count += 1
                    
                except Exception as e:
                    errors.append(f'{template_name}: {str(e)}')
            
            if processed_count == 0:
                return False, f'처리된 템플릿이 없습니다. 오류: {"; ".join(errors)}'
            
            # DOCX 파일 저장
            docx_path = self.output_folder / f"{year}년_{quarter}분기_지역경제동향_보도자료_전체.docx"
            doc.save(str(docx_path))
            
            # 결과 반환
            result = {
                'success': True,
                'output_filename': docx_path.name,
                'output_path': str(docx_path),
                'message': f'{processed_count}개 템플릿이 성공적으로 DOCX로 생성되었습니다.',
                'processed_templates': processed_count,
                'total_templates': len(TEMPLATE_ORDER)
            }
            
            if errors:
                result['warnings'] = errors
            
            return True, result
            
        except Exception as e:
            return False, f'서버 오류가 발생했습니다: {str(e)}'
        finally:
            excel_extractor.close()
