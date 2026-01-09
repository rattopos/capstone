"""
PDF를 읽어서 Word 템플릿으로 변환하는 모듈
PDF의 모든 페이지를 읽고 텍스트와 디자인을 그대로 재현한 Word 파일 생성
(표와 이미지는 제외하고 텍스트만 추출)
"""

import os
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from PIL import Image
import io

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

try:
    import cv2
    import numpy as np
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False


class PDFToWordConverter:
    """PDF를 읽어서 Word 템플릿으로 변환하는 클래스"""
    
    def __init__(self, use_easyocr: bool = True, dpi: int = 300):
        """
        PDF 변환기 초기화
        
        Args:
            use_easyocr: True면 easyocr 사용, False면 pytesseract 사용
            dpi: PDF를 이미지로 변환할 때의 DPI (기본값: 300)
        """
        self.use_easyocr = use_easyocr and EASYOCR_AVAILABLE
        self.dpi = dpi
        self.reader = None
        
        if self.use_easyocr:
            try:
                self.reader = easyocr.Reader(['ko', 'en'], gpu=False)
            except Exception as e:
                print(f"EasyOCR 초기화 실패, pytesseract 사용: {e}")
                self.use_easyocr = False
                self.reader = None
    
    def extract_text_from_pdf_direct(self, pdf_path: str) -> List[List[Dict]]:
        """
        PDF에서 직접 텍스트와 레이아웃 정보를 추출합니다 (이미지 변환 없이).
        
        Args:
            pdf_path: PDF 파일 경로
            
        Returns:
            페이지별 텍스트 정보 리스트. 각 페이지는 텍스트 정보 딕셔너리 리스트를 포함:
            - 'text': 추출된 텍스트
            - 'bbox': 바운딩 박스 좌표 (x1, y1, x2, y2)
            - 'font_size': 폰트 크기
            - 'is_bold': 굵은 글씨 여부
            - 'page_num': 페이지 번호
        """
        if not PYMUPDF_AVAILABLE:
            return []
        
        pages_text_data = []
        
        try:
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text_data = []
                
                
                # PDF에서 텍스트 딕셔너리 형식으로 추출
                text_dict = page.get_text("dict")
                
                # 텍스트 블록 처리
                for block in text_dict.get("blocks", []):
                    if "lines" not in block:
                        continue
                    
                    for line in block["lines"]:
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue
                            
                            # 바운딩 박스
                            bbox = span.get("bbox", [0, 0, 0, 0])
                            
                            # 폰트 정보
                            font_size = span.get("size", 12)
                            font_flags = span.get("flags", 0)
                            is_bold = (font_flags & 16) != 0  # bit 4는 bold
                            
                            page_text_data.append({
                                'text': text,
                                'bbox': tuple(bbox),
                                'font_size': font_size,
                                'is_bold': is_bold,
                                'confidence': 1.0,  # 직접 추출이므로 신뢰도 100%
                                'page_num': page_num + 1
                            })
                
                pages_text_data.append(page_text_data)
            
            doc.close()
            
        except Exception as e:
            print(f"PDF 직접 텍스트 추출 실패: {e}")
            return []
        
        return pages_text_data
    
    def pdf_to_images(self, pdf_path: str) -> List[Image.Image]:
        """
        PDF 파일을 이미지 리스트로 변환합니다 (스캔된 PDF용).
        
        Args:
            pdf_path: PDF 파일 경로
            
        Returns:
            PIL Image 객체 리스트
        """
        images = []
        
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(pdf_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    zoom = self.dpi / 72.0
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    images.append(img)
                doc.close()
                return images
            except Exception as e:
                print(f"PyMuPDF 변환 실패: {e}")
                raise ValueError(f"PDF를 이미지로 변환할 수 없습니다: {e}")
        
        raise ValueError("PDF를 이미지로 변환할 수 있는 라이브러리가 설치되어 있지 않습니다. PyMuPDF를 설치해주세요.")
    
    def extract_text_and_layout(self, image: Image.Image, ocr_progress_callback=None) -> List[Dict]:
        """
        이미지에서 텍스트와 레이아웃 정보를 추출합니다.
        
        Args:
            image: PIL Image 객체
            ocr_progress_callback: OCR 진행 상황 콜백 함수 (progress_percent: int) -> None
            
        Returns:
            텍스트 정보 딕셔너리 리스트. 각 딕셔너리는 다음 키를 포함:
            - 'text': 추출된 텍스트
            - 'bbox': 바운딩 박스 좌표 (x1, y1, x2, y2)
            - 'confidence': 신뢰도 (0-1)
            - 'font_size': 추정된 폰트 크기
            - 'is_bold': 굵은 글씨 여부 (추정)
        """
        text_data = []
        
        if CV2_AVAILABLE:
            img_array = np.array(image)
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_image = Image.fromarray(binary)
        else:
            processed_image = image.convert('L')
        
        # 이미지 전처리 완료 (10%)
        if ocr_progress_callback:
            ocr_progress_callback(10)
        
        if self.use_easyocr and self.reader:
            try:
                if CV2_AVAILABLE:
                    img_array = np.array(processed_image)
                    results = self.reader.readtext(img_array)
                else:
                    results = self.reader.readtext(processed_image)
                
                # OCR 결과 처리 시작 (20%)
                if ocr_progress_callback:
                    ocr_progress_callback(20)
                
                total_results = len(results)
                for idx, (bbox, text, confidence) in enumerate(results):
                    x_coords = [point[0] for point in bbox]
                    y_coords = [point[1] for point in bbox]
                    x1, y1 = min(x_coords), min(y_coords)
                    x2, y2 = max(x_coords), max(y_coords)
                    
                    # 폰트 크기 추정 (높이 기준)
                    font_size = max(8, int((y2 - y1) * 0.75))
                    
                    text_data.append({
                        'text': text.strip(),
                        'bbox': (x1, y1, x2, y2),
                        'confidence': confidence,
                        'font_size': font_size,
                        'is_bold': False  # OCR로는 정확히 판단 어려움
                    })
                    
                    # 진행 상황 업데이트 (20% ~ 90%)
                    if ocr_progress_callback and total_results > 0:
                        progress = 20 + int((idx + 1) / total_results * 70)
                        ocr_progress_callback(progress)
                
                # OCR 완료 (100%)
                if ocr_progress_callback:
                    ocr_progress_callback(100)
            except Exception as e:
                print(f"EasyOCR 처리 중 오류: {e}")
                if PYTESSERACT_AVAILABLE:
                    return self._extract_with_pytesseract(processed_image, ocr_progress_callback)
        
        elif PYTESSERACT_AVAILABLE:
            return self._extract_with_pytesseract(processed_image, ocr_progress_callback)
        else:
            raise ValueError("OCR 라이브러리가 설치되어 있지 않습니다. easyocr 또는 pytesseract를 설치해주세요.")
        
        return text_data
    
    def _extract_with_pytesseract(self, image: Image.Image, ocr_progress_callback=None) -> List[Dict]:
        """pytesseract를 사용하여 텍스트 추출"""
        text_data = []
        try:
            # OCR 시작 (30%)
            if ocr_progress_callback:
                ocr_progress_callback(30)
            
            data = pytesseract.image_to_data(image, lang='kor+eng', output_type=pytesseract.Output.DICT)
            
            # OCR 결과 처리 시작 (50%)
            if ocr_progress_callback:
                ocr_progress_callback(50)
            
            total_items = len(data['text'])
            for i in range(total_items):
                text = data['text'][i].strip()
                if text and int(data['conf'][i]) > 0:
                    x1 = data['left'][i]
                    y1 = data['top'][i]
                    x2 = x1 + data['width'][i]
                    y2 = y1 + data['height'][i]
                    confidence = float(data['conf'][i]) / 100.0
                    font_size = max(8, int(data['height'][i] * 0.75))
                    
                    text_data.append({
                        'text': text,
                        'bbox': (x1, y1, x2, y2),
                        'confidence': confidence,
                        'font_size': font_size,
                        'is_bold': False
                    })
                
                # 진행 상황 업데이트 (50% ~ 100%)
                if ocr_progress_callback and total_items > 0:
                    progress = 50 + int((i + 1) / total_items * 50)
                    ocr_progress_callback(progress)
            
            # OCR 완료 (100%)
            if ocr_progress_callback:
                ocr_progress_callback(100)
        except Exception as e:
            raise ValueError(f"OCR 처리 중 오류 발생: {e}")
        
        return text_data
    
    def organize_text_by_layout(self, text_data: List[Dict], page_width: int, page_height: int) -> List[Dict]:
        """
        텍스트 데이터를 레이아웃에 따라 정리합니다.
        
        Args:
            text_data: extract_text_and_layout()의 결과
            page_width: 페이지 너비
            page_height: 페이지 높이
            
        Returns:
            정리된 텍스트 구조 리스트 (단락으로 구분)
        """
        if not text_data:
            return []
        
        # y 좌표 기준으로 정렬
        sorted_text = sorted(text_data, key=lambda x: (x['bbox'][1], x['bbox'][0]))
        
        organized = []
        current_paragraph = []
        current_y = None
        line_threshold = page_height * 0.02
        
        for item in sorted_text:
            y = item['bbox'][1]
            
            if current_y is None or abs(y - current_y) < line_threshold:
                # 같은 줄
                current_paragraph.append(item)
                if current_y is None:
                    current_y = y
            else:
                # 새 줄
                if current_paragraph:
                    current_paragraph.sort(key=lambda x: x['bbox'][0])
                    organized.append({
                        'type': 'paragraph',
                        'items': current_paragraph,
                        'y': current_y,
                        'font_size': max([item['font_size'] for item in current_paragraph], default=12)
                    })
                
                current_paragraph = [item]
                current_y = y
        
        # 마지막 단락 처리
        if current_paragraph:
            current_paragraph.sort(key=lambda x: x['bbox'][0])
            organized.append({
                'type': 'paragraph',
                'items': current_paragraph,
                'y': current_y,
                'font_size': max([item['font_size'] for item in current_paragraph], default=12)
            })
        
        # y 좌표 기준으로 정렬
        organized.sort(key=lambda x: x.get('y', 0))
        
        return organized
    
    def convert_pdf_to_word(self, pdf_path: str, output_path: str, progress_callback=None, ocr_progress_callback=None) -> str:
        """
        PDF 파일을 Word 템플릿으로 변환합니다.
        모든 PDF를 이미지로 변환한 후 OCR을 사용하여 텍스트를 추출합니다.
        
        Args:
            pdf_path: PDF 파일 경로
            output_path: 출력 Word 파일 경로
            progress_callback: 진행 상황 콜백 함수 (current_page, total_pages, message, ocr_progress=None) -> None
            ocr_progress_callback: OCR 진행 상황 콜백 함수 (ocr_progress: int) -> None
            
        Returns:
            생성된 Word 파일 경로
        """
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx가 설치되어 있지 않습니다. pip install python-docx를 실행해주세요.")
        
        if not PYMUPDF_AVAILABLE:
            raise ValueError("PyMuPDF가 설치되어 있지 않습니다. pip install PyMuPDF를 실행해주세요.")
        
        print(f"PDF 파일 로딩 중: {pdf_path}")
        print("[DEBUG PDFToWord] 모든 PDF를 이미지로 변환한 후 OCR을 사용합니다.")
        
        # PDF 열기 (페이지 크기 정보를 위해)
        pdf_doc = fitz.open(pdf_path)
        total_pages = len(pdf_doc)
        
        # 첫 페이지의 크기로 Word 문서 설정
        first_page = pdf_doc[0]
        first_page_rect = first_page.rect
        pdf_page_width_pt = first_page_rect.width
        pdf_page_height_pt = first_page_rect.height
        
        # 포인트를 인치로 변환 (72 DPI 기준)
        pdf_page_width_inch = pdf_page_width_pt / 72.0
        pdf_page_height_inch = pdf_page_height_pt / 72.0
        
        # Word 문서 생성
        doc = Document()
        
        # 페이지 설정 (PDF 페이지 크기에 맞춤)
        section = doc.sections[0]
        section.page_height = Inches(pdf_page_height_inch)
        section.page_width = Inches(pdf_page_width_inch)
        
        # 마진 설정 (PDF와 유사하게)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        print(f"[DEBUG PDFToWord] PDF 페이지 크기: {pdf_page_width_inch:.2f} x {pdf_page_height_inch:.2f} 인치")
        
        # 각 페이지 처리
        for page_num in range(total_pages):
            print(f"페이지 {page_num + 1}/{total_pages} 처리 중...")
            
            # 진행 상황 콜백 호출
            if progress_callback:
                progress_callback(page_num + 1, total_pages, f"페이지 {page_num + 1}/{total_pages} 처리 중...")
            
            if page_num > 0:
                # 페이지 나누기
                doc.add_page_break()
            
            page = pdf_doc[page_num]
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # 모든 페이지를 이미지로 변환 후 OCR 사용
            zoom = self.dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            
            # OCR로 텍스트 추출
            print(f"  - OCR로 텍스트 추출 중...")
            
            # OCR 진행 상황 콜백 래퍼 함수
            def ocr_progress_wrapper(ocr_progress):
                """OCR 진행률을 전체 진행률에 반영"""
                if progress_callback:
                    # 페이지 내 OCR 진행률을 전체 진행률에 반영
                    # 1단계의 40% 중에서 현재 페이지의 OCR 진행률 반영
                    page_base_progress = int((page_num / total_pages) * 40)
                    page_ocr_progress = int((ocr_progress / 100) * (40 / total_pages))
                    overall_progress = 5 + page_base_progress + page_ocr_progress
                    
                    message = f"페이지 {page_num + 1}/{total_pages}: OCR 처리 중"
                    if ocr_progress < 100:
                        message += f" ({ocr_progress}%)"
                    
                    progress_callback(page_num + 1, total_pages, message, ocr_progress)
            
            # OCR 진행 상황 콜백 전달
            effective_ocr_callback = ocr_progress_callback if ocr_progress_callback else ocr_progress_wrapper
            
            text_data = self.extract_text_and_layout(image, ocr_progress_callback=effective_ocr_callback)
            organized_text = self.organize_text_by_layout(text_data, image.size[0], image.size[1]) if text_data else []
            
            if not organized_text:
                # 텍스트가 없으면 빈 단락 추가
                print(f"[DEBUG PDFToWord] ⚠️ 경고: 페이지 {page_num + 1}에서 텍스트를 추출할 수 없습니다")
                para = doc.add_paragraph()
                para.add_run(f"(페이지 {page_num + 1} - 텍스트를 추출할 수 없습니다)")
                continue
            
            # 디버그: 첫 페이지의 텍스트 샘플 출력
            if page_num == 0:
                print(f"[DEBUG PDFToWord] 첫 페이지 텍스트 샘플 (처음 3개 요소):")
                for i, item in enumerate(organized_text[:3]):
                    item_type = item.get('type', 'unknown')
                    if item_type == 'paragraph' and 'items' in item:
                        text_sample = ' '.join([ti.get('text', '') for ti in item.get('items', [])])
                        print(f"  요소 {i+1} (단락): {repr(text_sample[:100])}")
                    else:
                        print(f"  요소 {i+1} (타입: {item_type})")
            
            # 텍스트 요소를 y 좌표 기준으로 정렬
            all_elements = []
            
            # 텍스트 요소 추가
            for item in organized_text:
                # 안전하게 타입과 y 좌표 가져오기
                item_type = item.get('type', 'paragraph')
                item_y = item.get('y', 0)
                
                all_elements.append({
                    'type': item_type,
                    'data': item,
                    'y': item_y
                })
            
            # y 좌표 기준으로 정렬
            all_elements.sort(key=lambda x: x.get('y', 0))
            
            # Word 문서에 추가
            prev_y = None
            for element in all_elements:
                try:
                    element_y = element.get('y', 0)
                    
                    # 큰 간격이 있으면 단락 구분
                    if prev_y is not None and (element_y - prev_y) > page_height * 0.05:
                        doc.add_paragraph()  # 빈 단락 추가
                    
                    # 단락인 경우
                    if element.get('type') == 'paragraph':
                        item = element.get('data', {})
                        if not item:
                            continue
                        
                        para = doc.add_paragraph()
                        
                        # 텍스트 아이템들을 순서대로 추가 (폰트 정보 유지)
                        items = item.get('items', [])
                        if items and isinstance(items, list):
                            for idx, text_item in enumerate(items):
                                if not isinstance(text_item, dict):
                                    continue
                                
                                text = text_item.get('text', '')
                                if not text:
                                    continue
                                
                                # 숫자나 데이터를 의미 기반 마커로 변환 (단일 아이템 기준)
                                processed_text = self._convert_data_to_semantic_markers(text, [text_item])
                                
                                # 각 텍스트 아이템을 별도 Run으로 추가 (폰트 정보 유지)
                                run = para.add_run(processed_text)
                                
                                # 폰트 크기 설정 (정확한 크기 유지)
                                font_size = text_item.get('font_size', 12)
                                if font_size:
                                    run.font.size = Pt(font_size)
                                
                                # 폰트 이름 설정 (한글 지원)
                                run.font.name = '맑은 고딕'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                                
                                # 굵은 글씨 설정
                                if text_item.get('is_bold', False):
                                    run.bold = True
                                
                                # 다음 아이템과 공백 추가 (같은 줄의 경우)
                                if idx < len(items) - 1 and isinstance(items[idx + 1], dict):
                                    # 다음 아이템과의 x 좌표 차이 확인
                                    current_bbox = text_item.get('bbox', [0, 0, 0, 0])
                                    next_bbox = items[idx + 1].get('bbox', [0, 0, 0, 0])
                                    if len(current_bbox) >= 3 and len(next_bbox) >= 1:
                                        current_x = current_bbox[2]
                                        next_x = next_bbox[0]
                                        gap = next_x - current_x
                                        
                                        # 적절한 간격이면 공백 추가
                                        if 5 < gap < 100:  # 5~100 픽셀 간격
                                            run.add_text(' ')
                        else:
                            # 아이템이 없는 경우 빈 단락
                            para.add_run('')
                        
                        # 단락 정렬 (왼쪽 정렬 기본)
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # 줄 간격 설정 (원본과 유사하게)
                        para.paragraph_format.line_spacing = 1.15
                        
                        # 단락 간격 설정 (y 좌표 차이 기반)
                        if prev_y is not None:
                            y_diff = element_y - prev_y
                            # 큰 간격이면 단락 간격 추가
                            if y_diff > page_height * 0.02:
                                para.paragraph_format.space_before = Pt(max(6, int(y_diff * 0.1)))
                        
                        prev_y = element_y
                except Exception as e:
                    print(f"[DEBUG PDFToWord] 요소 처리 중 오류 발생: {e}")
                    import traceback
                    traceback.print_exc()
                    # 오류가 발생해도 계속 진행
                    continue
        
        pdf_doc.close()
        
        # Word 파일 저장
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))
        print(f"Word 템플릿이 생성되었습니다: {output_path}")
        
        return str(output_file)
    
    def _convert_data_to_semantic_markers(self, text: str, text_items: List[Dict]) -> str:
        """
        텍스트에서 숫자나 데이터를 의미 기반 마커로 변환합니다.
        하드코딩된 시트명 대신 의미 기반 키워드를 사용합니다.
        
        Args:
            text: 원본 텍스트
            text_items: 텍스트 아이템 리스트 (컨텍스트 분석용)
            
        Returns:
            마커가 포함된 텍스트
        """
        import re
        
        # 숫자 패턴 (퍼센트, 소수점, 콤마 포함)
        number_pattern = r'[\d,]+\.?\d*%?'
        
        # 숫자가 포함된 경우 주변 컨텍스트 분석
        def replace_with_marker(match):
            number = match.group(0)
            
            # 주변 텍스트에서 키워드 추출
            start_pos = match.start()
            end_pos = match.end()
            
            # 앞뒤 20자 정도의 컨텍스트 추출
            context_start = max(0, start_pos - 20)
            context_end = min(len(text), end_pos + 20)
            context = text[context_start:context_end]
            
            # 의미 기반 키워드 추출
            sheet_keyword = self._extract_sheet_keyword_from_context(context)
            data_keyword = self._extract_data_keyword_from_context(context, number)
            
            # 마커 생성: {의미키워드:데이터키워드}
            if sheet_keyword and data_keyword:
                return f"{{{sheet_keyword}:{data_keyword}}}"
            elif sheet_keyword:
                return f"{{{sheet_keyword}:값}}"
            else:
                # 키워드를 추출하지 못한 경우 원본 숫자 유지
                return number
        
        # 숫자 패턴을 마커로 변환
        result = re.sub(number_pattern, replace_with_marker, text)
        
        return result
    
    def _extract_sheet_keyword_from_context(self, context: str) -> Optional[str]:
        """컨텍스트에서 시트 키워드 추출"""
        import re
        
        # 키워드 매핑 (의미 기반)
        keyword_patterns = {
            '경제지표': r'경제|지표|gdp|생산|소비|투자',
            '생산': r'생산|제조|manufacturing|production',
            '소비': r'소비|소매|consumption|retail',
            '건설': r'건설|construction|공정',
            '광업': r'광업|mining',
            '제조업': r'제조|manufacturing',
            '서비스업': r'서비스|service',
            '소매업': r'소매|retail',
            '지역': r'지역|region|시도|시군구',
            '전국': r'전국|national|전체',
        }
        
        context_lower = context.lower()
        
        for keyword, pattern in keyword_patterns.items():
            if re.search(pattern, context_lower, re.IGNORECASE):
                return keyword
        
        return None
    
    def _extract_data_keyword_from_context(self, context: str, number: str) -> Optional[str]:
        """컨텍스트에서 데이터 키워드 추출"""
        import re
        
        # 데이터 키워드 패턴
        data_patterns = {
            '증감률': r'증감|증가|감소|변화|변동|growth|change',
            '증가율': r'증가|상승|up|rise',
            '감소율': r'감소|하락|down|fall',
            '전국': r'전국|national|전체',
            '서울': r'서울|seoul',
            '부산': r'부산|busan',
            '대구': r'대구|daegu',
            '인천': r'인천|incheon',
            '광주': r'광주|gwangju',
            '대전': r'대전|daejeon',
            '울산': r'울산|ulsan',
        }
        
        context_lower = context.lower()
        
        keywords = []
        for keyword, pattern in data_patterns.items():
            if re.search(pattern, context_lower, re.IGNORECASE):
                keywords.append(keyword)
        
        if keywords:
            # 여러 키워드가 있으면 언더스코어로 연결
            return '_'.join(keywords)
        
        # 키워드를 찾지 못한 경우 숫자 주변 텍스트에서 추출
        # 예: "전국 2.5%" -> "전국_값"
        words = re.findall(r'[가-힣]+', context)
        if words:
            return '_'.join(words[:2])  # 최대 2개 단어
        
        return None

