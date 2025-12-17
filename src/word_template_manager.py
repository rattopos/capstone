"""
Word 템플릿 관리 모듈
Word 템플릿 파일 로드 및 마커 파싱 기능 제공
"""

import re
from pathlib import Path
from typing import List, Dict, Tuple

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordTemplateManager:
    """Word 템플릿을 관리하고 마커를 파싱하는 클래스"""
    
    # 마커 패턴: 
    # 1. {시트키워드:데이터키워드} - 의미 기반 마커
    # 2. {데이터키워드} - 시트 키워드 생략 가능
    # 3. {시트명:셀주소} - 기존 형식 (하위 호환)
    # 4. {시트명:셀주소:계산식} - 계산식 포함
    MARKER_PATTERN = re.compile(r'\{([^:{}]*):?([^:}]+)(?::([^}]+))?\}')
    
    def __init__(self, template_path: str):
        """
        Word 템플릿 매니저 초기화
        
        Args:
            template_path: Word 템플릿 파일 경로 (.docx)
        """
        self.template_path = Path(template_path)
        self.document = None
        self.markers = []
        
    def load_template(self) -> Document:
        """
        Word 템플릿 파일을 로드합니다.
        
        Returns:
            Document 객체
            
        Raises:
            FileNotFoundError: 템플릿 파일이 존재하지 않을 때
            IOError: 파일 읽기 실패 시
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx가 설치되어 있지 않습니다. pip install python-docx를 실행해주세요.")
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {self.template_path}")
        
        try:
            self.document = Document(str(self.template_path))
            return self.document
        except IOError as e:
            raise IOError(f"템플릿 파일 읽기 실패: {e}")
    
    def extract_markers(self) -> List[Dict[str, any]]:
        """
        Word 템플릿에서 모든 마커를 추출합니다.
        
        Returns:
            마커 정보 딕셔너리 리스트. 각 딕셔너리는 다음 키를 포함:
            - 'full_match': 전체 마커 문자열 (예: '{시트1:A1:sum}')
            - 'sheet_name': 시트명
            - 'cell_address': 셀 주소 (또는 범위 또는 동적 키)
            - 'operation': 계산식 (선택적, 없으면 None)
            - 'paragraph': 마커가 포함된 Paragraph 객체
            - 'run': 마커가 포함된 Run 객체
            - 'run_index': Run의 인덱스
        """
        if self.document is None:
            self.load_template()
        
        self.markers = []
        
        print(f"[DEBUG WordTemplateManager] 단락 개수: {len(self.document.paragraphs)}")
        print(f"[DEBUG WordTemplateManager] 테이블 개수: {len(self.document.tables)}")
        
        # 모든 단락에서 마커 찾기
        for para_idx, paragraph in enumerate(self.document.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                # 마커 패턴이 있는지 미리 확인
                if self.MARKER_PATTERN.search(para_text):
                    print(f"[DEBUG WordTemplateManager] 단락 {para_idx+1}에서 마커 패턴 발견: {repr(para_text[:100])}")
            markers_in_para = self._extract_markers_from_paragraph(paragraph, para_idx)
            if markers_in_para:
                print(f"[DEBUG WordTemplateManager] 단락 {para_idx+1}에서 {len(markers_in_para)}개 마커 추출")
            self.markers.extend(markers_in_para)
        
        # 모든 테이블에서 마커 찾기
        for table_idx, table in enumerate(self.document.tables):
            markers_in_table = self._extract_markers_from_table(table, table_idx)
            if markers_in_table:
                print(f"[DEBUG WordTemplateManager] 테이블 {table_idx+1}에서 {len(markers_in_table)}개 마커 추출")
            self.markers.extend(markers_in_table)
        
        print(f"[DEBUG WordTemplateManager] 총 {len(self.markers)}개 마커 추출 완료")
        
        return self.markers
    
    def _extract_markers_from_paragraph(self, paragraph: Paragraph, para_idx: int) -> List[Dict[str, any]]:
        """단락에서 마커 추출"""
        markers = []
        full_text = paragraph.text
        
        # 마커 패턴 찾기
        matches = self.MARKER_PATTERN.finditer(full_text)
        
        for match in matches:
            sheet_part = match.group(1).strip() if match.group(1) else None
            data_part = match.group(2).strip()
            operation = match.group(3).strip() if match.group(3) else None
            
            # 마커가 포함된 Run 찾기
            run_info = self._find_run_containing_marker(paragraph, match.start(), match.end())
            
            # 마커 타입 판단 (셀 주소 형식인지 의미 기반인지)
            is_cell_address = bool(re.match(r'^[A-Z]+\d+', data_part))
            
            marker_info = {
                'full_match': match.group(0),
                'sheet_keyword': sheet_part,  # 시트 키워드 (None일 수 있음)
                'data_keyword': data_part,    # 데이터 키워드 또는 셀 주소
                'sheet_name': sheet_part,     # 하위 호환성
                'cell_address': data_part,    # 하위 호환성
                'operation': operation,
                'is_semantic': not is_cell_address,  # 의미 기반 마커인지
                'paragraph': paragraph,
                'run': run_info['run'] if run_info else None,
                'run_index': run_info['index'] if run_info else None,
                'para_index': para_idx,
                'start_pos': match.start(),
                'end_pos': match.end()
            }
            
            markers.append(marker_info)
        
        return markers
    
    def _extract_markers_from_table(self, table: Table, table_idx: int) -> List[Dict[str, any]]:
        """테이블에서 마커 추출"""
        markers = []
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    full_text = paragraph.text
                    matches = self.MARKER_PATTERN.finditer(full_text)
                    
                    for match in matches:
                        sheet_part = match.group(1).strip() if match.group(1) else None
                        data_part = match.group(2).strip()
                        operation = match.group(3).strip() if match.group(3) else None
                        
                        run_info = self._find_run_containing_marker(paragraph, match.start(), match.end())
                        
                        # 마커 타입 판단
                        is_cell_address = bool(re.match(r'^[A-Z]+\d+', data_part))
                        
                        marker_info = {
                            'full_match': match.group(0),
                            'sheet_keyword': sheet_part,
                            'data_keyword': data_part,
                            'sheet_name': sheet_part,     # 하위 호환성
                            'cell_address': data_part,    # 하위 호환성
                            'operation': operation,
                            'is_semantic': not is_cell_address,
                            'paragraph': paragraph,
                            'run': run_info['run'] if run_info else None,
                            'run_index': run_info['index'] if run_info else None,
                            'table_index': table_idx,
                            'row_index': row_idx,
                            'col_index': col_idx,
                            'para_index': para_idx,
                            'start_pos': match.start(),
                            'end_pos': match.end()
                        }
                        
                        markers.append(marker_info)
        
        return markers
    
    def _find_run_containing_marker(self, paragraph: Paragraph, start_pos: int, end_pos: int) -> Dict:
        """마커가 포함된 Run 찾기"""
        current_pos = 0
        
        for run_idx, run in enumerate(paragraph.runs):
            run_text = run.text
            run_start = current_pos
            run_end = current_pos + len(run_text)
            
            # 마커가 이 Run에 포함되어 있는지 확인
            if run_start <= start_pos < run_end or run_start < end_pos <= run_end:
                return {
                    'run': run,
                    'index': run_idx,
                    'start_offset': start_pos - run_start,
                    'end_offset': end_pos - run_start
                }
            
            current_pos = run_end
        
        return None
    
    def replace_marker(self, marker_info: Dict, value: str) -> None:
        """
        Word 템플릿에서 특정 마커를 값으로 치환합니다.
        
        Args:
            marker_info: extract_markers()에서 반환된 마커 정보
            value: 치환할 값
        """
        paragraph = marker_info['paragraph']
        full_match = marker_info['full_match']
        start_pos = marker_info['start_pos']
        end_pos = marker_info['end_pos']
        
        # 단락의 전체 텍스트에서 마커를 값으로 치환
        full_text = paragraph.text
        new_text = full_text[:start_pos] + str(value) + full_text[end_pos:]
        
        # 단락의 모든 Run 제거 후 새로 생성
        paragraph.clear()
        
        # 새 Run 생성 (기존 Run의 스타일 유지)
        run = paragraph.add_run(new_text)
        if marker_info.get('run'):
            original_run = marker_info['run']
            # 스타일 복사
            try:
                run.font.name = original_run.font.name
            except:
                pass
            try:
                run.font.size = original_run.font.size
            except:
                pass
            try:
                run.font.bold = original_run.font.bold
            except:
                pass
            try:
                run.font.italic = original_run.font.italic
            except:
                pass
            try:
                run.font.underline = original_run.font.underline
            except:
                pass
            try:
                if hasattr(original_run.font, 'color') and original_run.font.color:
                    run.font.color = original_run.font.color
            except:
                pass
    
    def get_template_content(self) -> Document:
        """
        현재 로드된 Word 템플릿을 반환합니다.
        
        Returns:
            Document 객체
        """
        if self.document is None:
            self.load_template()
        return self.document
    
    def save_template(self, output_path: str) -> None:
        """
        템플릿을 파일로 저장합니다.
        
        Args:
            output_path: 저장할 파일 경로
        """
        if self.document is None:
            raise ValueError("템플릿이 로드되지 않았습니다. load_template()를 먼저 호출하세요.")
        
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_file))
    
    def validate_template(self, required_markers: List[str] = None) -> Tuple[bool, List[str]]:
        """
        템플릿의 유효성을 검증합니다.
        
        Args:
            required_markers: 필수 마커 리스트 (전체 마커 문자열 형식)
            
        Returns:
            (유효성 여부, 누락된 마커 리스트) 튜플
        """
        if self.document is None:
            self.load_template()
        
        extracted = self.extract_markers()
        extracted_full_matches = [m['full_match'] for m in extracted]
        
        if required_markers is None:
            return True, []
        
        missing_markers = [
            marker for marker in required_markers 
            if marker not in extracted_full_matches
        ]
        
        is_valid = len(missing_markers) == 0
        return is_valid, missing_markers

